"""常见文档文本导入工具。"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


SUPPORTED_DOCUMENT_SUFFIXES = {".pdf", ".docx", ".md", ".markdown", ".txt"}
MAX_IMPORT_FILES = 20
MAX_IMPORT_FILE_BYTES = 10 * 1024 * 1024
MAX_IMPORT_TOTAL_CHARS = 80000
DEFAULT_CHUNK_CHARS = 900
DEFAULT_CHUNK_OVERLAP = 120


@dataclass(frozen=True)
class ImportedDocument:
    path: str
    title: str
    text: str


def import_documents(paths: list[str | Path], max_total_chars: int = MAX_IMPORT_TOTAL_CHARS) -> list[ImportedDocument]:
    """从常见文档中提取文本，限制文件数量、大小和总字符数。"""

    if len(paths) > MAX_IMPORT_FILES:
        raise ValueError(f"一次最多导入 {MAX_IMPORT_FILES} 个文档")

    imported: list[ImportedDocument] = []
    remaining = max_total_chars
    for raw_path in paths:
        path = Path(raw_path)
        if remaining <= 0:
            break
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_DOCUMENT_SUFFIXES:
            raise ValueError(f"不支持的文档类型: {path.name}")
        if path.stat().st_size > MAX_IMPORT_FILE_BYTES:
            raise ValueError(f"文件过大: {path.name}，单个文件最多 10MB")

        text = _read_document_text(path)
        text = _normalize_text(text)
        if not text:
            continue
        clipped = text[:remaining]
        imported.append(ImportedDocument(str(path), path.name, clipped))
        remaining -= len(clipped)
    return imported


def render_imported_documents(documents: list[ImportedDocument]) -> str:
    """渲染为可追加到事件背景中的文本块。"""

    parts: list[str] = []
    for document in documents:
        parts.append(f"【导入文档：{document.title}】\n{document.text}")
    return "\n\n".join(parts)


def chunk_text(
    text: str,
    *,
    max_chars: int = DEFAULT_CHUNK_CHARS,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> list[str]:
    """将长文本切成带少量重叠的 RAG 片段。"""

    normalized = _normalize_text(text)
    if not normalized:
        return []
    chunks: list[str] = []
    start = 0
    safe_overlap = min(overlap, max_chars // 3)
    while start < len(normalized):
        end = min(start + max_chars, len(normalized))
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start = max(end - safe_overlap, start + 1)
    return chunks


def _read_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # 中文 Windows 的 ANSI(GBK) 文本：回退 GBK，忽略仍无法解码的字符
            return path.read_text(encoding="gbk", errors="ignore")
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("导入 PDF 需要安装 pypdf：pip install pypdf") from exc
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("导入 Word 文档需要安装 python-docx：pip install python-docx") from exc
        doc = Document(str(path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    raise ValueError(f"不支持的文档类型: {path.name}")


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compact: list[str] = []
    previous_blank = False
    for line in lines:
        is_blank = not line
        if is_blank and previous_blank:
            continue
        compact.append(line)
        previous_blank = is_blank
    return "\n".join(compact).strip()
